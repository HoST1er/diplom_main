import json
from collections import defaultdict

import docx
from django.core.exceptions import PermissionDenied
from docx import Document

import openpyxl
from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from django.core.files import temp
from django.core.files.storage import FileSystemStorage, default_storage
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, redirect
from django.urls import reverse_lazy
from django.views.decorators.csrf import csrf_exempt, csrf_protect
from django.views.generic import FormView
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

from diplommain import settings
from .models import FirstVariantBd, Description_of_competencies, UserSessionData
from main.scripts.parse_excel import Command
from django.contrib import messages
import fnmatch
import os
from .forms import FileUploadForm, RegisterForm
from django.contrib.sessions.models import Session
from django.contrib.auth.models import User
from django.utils import timezone
import pickle
from django.conf import settings
from importlib import import_module


#user_data = {}

def debug_sessions_view(request):
    sessions = Session.objects.filter(expire_date__gte=timezone.now())
    print(f'--- Найдено {sessions.count()} сессий ---')

    for session in sessions:
        data = session.get_decoded()
        print('Сессия:', session.session_key)
        print('Пользователь ID:', data.get('_auth_user_id'))
        print('user_data:', data.get('user_data'))
        print('---')

    return HttpResponse("Смотри консоль сервера.")

@login_required
def finish_work(request):
    if 'user_data' in request.session:
        UserSessionData.objects.create(
            user=request.user,
            data=request.session['user_data']
        )
        del request.session['user_data']  # можно очистить, если нужно
    return redirect('home')  # или куда нужно




@login_required
def profile_view(request):
    if request.user.is_superuser:
        sessions = Session.objects.filter(expire_date__gte=timezone.now())
        session_data_list = []

        # Для доступа к session store (внутренности сессии)
        engine = import_module(settings.SESSION_ENGINE)

        for session in sessions:
            try:
                data = session.get_decoded()
                print('Decoded data:', data)  # <-- добавь это
                user_id = data.get('_auth_user_id')
                user = User.objects.get(id=user_id)

                if 'user_data' in data:
                    session_data_list.append({
                        'user': user,
                        'user_data': data['user_data'],
                        'source': 'active'
                    })
            except:
                continue

        completed = UserSessionData.objects.all()

        for item in completed:
            session_data_list.append({
                'user': item.user,
                'user_data': item.data,
                'source': 'saved'
            })

        return render(request, 'main/admin_session_data.html', {'session_data_list': session_data_list})
    else:
        if request.path == '/admin_session_data/':
            raise PermissionDenied()
        # Данные и шаблон для обычного пользователя
        return render(request, 'main/home.html', {'user': request.user})
def logout_view(request):
    if 'user_data' in request.session:
        print("✅ user_data найден в сессии. Сохраняем...")
        UserSessionData.objects.create(
            user=request.user,
            data=request.session['user_data']
        )
    else:
        print("⚠️ user_data отсутствует в сессии.")
    logout(request)
    return redirect('profile_view')

class RegisterView(FormView):
    form_class = RegisterForm
    template_name = 'registration/register.html'
    success_url = reverse_lazy('profile_view')

    def form_valid(self, form):
        form.save()
        return super().form_valid(form)



# def admin_user_data_view(request):
#     if not request.user.is_superuser:
#         return redirect('dashboard')  # или вернуть 403
#
#     sessions = Session.objects.filter(expire_date__gte=timezone.now())
#     session_data_list = []
#
#     # Для доступа к session store (внутренности сессии)
#     engine = import_module(settings.SESSION_ENGINE)
#
#     for session in sessions:
#         try:
#             data = session.get_decoded()
#             user_id = data.get('_auth_user_id')
#             user = User.objects.get(id=user_id)
#
#             if 'user_data' in data:
#                 session_data_list.append({
#                     'user': user,
#                     'user_data': data['user_data']
#                 })
#
#         except Exception as e:
#             continue  # если пользователь удалён или сессия битая
#
#     return render(request, 'main/admin_session_data.html', {'session_data_list': session_data_list})



@csrf_exempt  # Отключение CSRF-защиты (не рекомендуется для продакшена!)
def parse_excel_files(request):
    if request.method == 'POST' and request.FILES.getlist('files'):
        print("Файл получен:", request.FILES.getlist('files'))
        files = request.FILES.getlist('files')
        temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp_dir')
        os.makedirs(temp_dir, exist_ok=True)
        print(files)
        for file in files:
            print(file)
            # Сохранение файла во временное хранилище
            file_path = default_storage.save(os.path.join("temp_dir", file.name), file)
            #file_path = default_storage.save(f"/media/temp_dir/{file.name}", file)
            print(file_path)
            #try:
            try:
                wb = openpyxl.load_workbook(default_storage.path(file_path))
                sheet = wb['Дисциплины']

                if sheet['L17'].value == 'Расчетно-аналитическая работа':
                    print('есть столбец расчетно-аналитическая работа')
                    # Читаем профиль
                    try:
                        try:
                            cell_value = sheet['A11'].value.split(',')
                            direction_of_preparation = cell_value[0].split('направление подготовки: ')[1]
                            edu_program = cell_value[1].split(' ОП ')[1].replace('"', '')
                            profile = cell_value[2].split('Профиль: ')[1].replace('"', '')
                            scientific_speciality = ''
                        except:
                            cell_value = sheet['A12'].value.split(',')
                            direction_of_preparation = cell_value[0].split('направление подготовки: ')[1]
                            edu_program = cell_value[1].split('Направленность программы: ')[1].replace('"', '')
                            profile = ''
                            scientific_speciality = ''
                    except:
                        #print(f"Ошибка при обработке A11: {e}")
                        direction_of_preparation = ''
                        edu_program = ''
                        profile = ''

                        cell_value = sheet['A11'].value
                        cell_value = cell_value.split(': ')
                        #print(cell_value)

                        scientific_speciality = cell_value[1]
                            #print(scientific_speciality)

                    # Определяем колонки
                    columns = ['B', 'C', 'D', 'F', 'E', 'J',
                               'G', 'H', 'I', 'K', 'L', 'M',
                               'N', 'Q', 'R', 'S', 'T', 'O', 'P']
                    column_indices = [openpyxl.utils.column_index_from_string(col) for col in columns]
                    start_row = 23
                    data = []

                    # Читаем данные
                    for row in sheet.iter_rows(min_row=start_row, max_row=sheet.max_row):
                        row_data = []
                        is_bright = True

                        for idx in column_indices:
                            cell = row[idx - 1]
                            cell_color = cell.fill.start_color

                            if cell_color.index == '00000000' or cell_color.index == 'FFD5EFFF':
                                row_data.append(cell.value)
                            else:
                                is_bright = False
                                break

                        if is_bright:
                            data.append(row_data)

                    # Функции для преобразования данных
                    def to_int(value):
                        try:
                            return int(value) if value not in [None, '', ' '] else None
                        except ValueError:
                            return None

                    def to_float(value):
                        try:
                            return float(value) if value not in [None, '', ' '] else None
                        except ValueError:
                            return None

                    # Сохраняем в БД
                    for row in data:
                        print(row)
                        try:
                            print(f"Обрабатываем строку данных: {row}")
                            FirstVariantBd.objects.create(
                                name_object=row[0],
                                department=row[1].replace('\n', '') if row[1] else '',
                                competentions=row[2].replace('\n', '') if row[2] else '',
                                profile=profile,
                                direction_of_preparation=direction_of_preparation,
                                edu_program=edu_program,
                                test_obj=row[3] if row[3] else '',
                                exam=row[4] if row[4] else '',
                                control_work=row[5] if row[5] else '',
                                test_obj_with_mark=row[6] if row[6] else '',
                                course_work=row[7] if row[7] else '',
                                course_project=row[8] if row[8] else '',
                                essay=row[9] if row[9] else '',
                                calcul_analytic_work=row[10] if row[10] else '',
                                creative_homework=row[11] if row[11] else '',
                                project_work=row[12] if row[12] else '',
                                classroom_hours=to_int(row[13]),
                                lectures=to_int(row[14]),
                                seminars=to_int(row[15]),
                                independent_work=to_int(row[16]),
                                ECTS=to_float(row[17]),
                                total_hours=to_int(row[18]),
                                scientific_speciality=scientific_speciality
                            )
                        except Exception as e:
                            print(f"Ошибка обработки файла {os.path.basename(file_path)}: {e}")

                    # Удаление файла после обработки
                    default_storage.delete(file_path)

                else:
                    print('нет столбца расчетно-аналитическая работа')

                    try:
                        # Обработка данных из A11
                        try:
                            cell_value = sheet['A11'].value.split(',')
                            direction_of_preparation = cell_value[0].split('направление подготовки: ')[1]
                            edu_program = cell_value[1].split(' ОП ')[1].replace('"', '')
                            profile = cell_value[2].split('Профиль: ')[1].replace('"', '')
                            scientific_speciality = ''
                        except:
                            cell_value = sheet['A12'].value.split(',')
                            direction_of_preparation = cell_value[0].split('направление подготовки: ')[1]
                            edu_program = cell_value[1].split('Направленность программы: ')[1].replace('"', '')
                            profile = ''
                            scientific_speciality = ''
                    except:
                        direction_of_preparation, edu_program, profile = '', '', ''
                        scientific_speciality = sheet['A11'].value.split(': ')[1]

                    # Определение колонок для чтения
                    columns = ['B', 'C', 'D', 'F', 'E', 'J',
                               'G', 'H', 'I', 'K', 'L', 'M',
                               'N', 'Q', 'R', 'S', 'O', 'P']
                    column_indices = [openpyxl.utils.column_index_from_string(col) for col in columns]
                    start_row = 23

                    data = []
                    for row in sheet.iter_rows(min_row=start_row, max_row=sheet.max_row):
                        row_data = []
                        is_valid = True
                        for idx in column_indices:
                            cell = row[idx - 1]
                            cell_color = cell.fill.start_color.index
                            if cell_color not in ['00000000', 'FFD5EFFF']:  # Проверка цвета
                                is_valid = False
                                break
                            row_data.append(cell.value)
                        if is_valid:
                            data.append(row_data)

                    # Функции для приведения типов
                    def to_int(value):
                        try:
                            return int(value) if value else None
                        except ValueError:
                            return None

                    def to_float(value):
                        try:
                            return float(value) if value else None
                        except ValueError:
                            return None

                    # Запись данных в базу
                    for row in data:
                        try:
                            FirstVariantBd.objects.create(
                                name_object=row[0],
                                department=row[1].replace('\n', '') if row[1] else '',
                                competentions=row[2].replace('\n', '') if row[2] else '',
                                profile=profile,
                                direction_of_preparation=direction_of_preparation,
                                edu_program=edu_program,
                                test_obj=row[3] or '',
                                exam=row[4] or '',
                                control_work=row[5] or '',
                                test_obj_with_mark=row[6] or '',
                                course_work=row[7] or '',
                                course_project=row[8] or '',
                                essay=row[9] or '',
                                calcul_analytic_work='',
                                creative_homework=row[10] or '',
                                project_work=row[11] or '',
                                classroom_hours=to_int(row[17]),
                                lectures=to_int(row[13]),
                                seminars=to_int(row[14]),
                                independent_work=to_int(row[15]),
                                ECTS=to_float(row[12]),
                                total_hours=to_int(row[16]),
                                scientific_speciality=scientific_speciality
                            )
                        except Exception as e:
                            print(f"Ошибка обработки строки {row}: {e}")

                    # Удаление файла после обработки
                    default_storage.delete(file_path)
            except:
                default_storage.delete(file_path)
                continue

        return redirect("home")


            # except:
            #     try:
            #         # В случае ошибки обработки файла
            #         os.remove(file_path)  # Убедитесь, что файл удаляется при ошибке
            #         return JsonResponse({'status': 'error', 'message': f'Ошибка при обработке файла: {str(e)}'})
            #     except Exception as e:
            #         print(f"Ошибка при удалении временного файла {file_path}: {e}")

            # return render(request, 'success.html')

    else:
        form = FileUploadForm()
        return render(request, 'main/upload.html', {'form': form})

@csrf_exempt
def parse_competencies(request):
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']
        print("Файл получен:", uploaded_file.name)

        temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp_dir')
        os.makedirs(temp_dir, exist_ok=True)

        file_path = default_storage.save(os.path.join("temp_dir", uploaded_file.name), uploaded_file)
        print("Файл сохранен по пути:", file_path)

        wb = openpyxl.load_workbook(default_storage.path(file_path))

        sheet = wb.active  # Можно указать конкретный лист: wb['Sheet1']

        columns = ['A', 'B']

        column_indices = [openpyxl.utils.column_index_from_string(col) for col in columns]

        start_row = 2  # Например, начиная с третьей строки (индекс 3)

        data = []

        for row in sheet.iter_rows(min_row=start_row, max_row=sheet.max_row):
            row_data = []
            for idx in column_indices:
                cell = row[idx - 1]  # Индексы столбцов в openpyxl начинаются с 1, но Python использует 0-based индекс
                # print(cell.value)
                row_data.append(cell.value)
            data.append(row_data)

        for row in data:
            Description_of_competencies.objects.create(
                competency_name=row[0],
                description=row[1]
            )

        # Удаление файла после обработки
        default_storage.delete(file_path)

        return redirect("home")
    else:
        form = FileUploadForm()
        return render(request, 'main/upload_competencies.html', {'form': form})

def home(request):
    main_bd = FirstVariantBd.objects.all()[:1]
    return render(request, 'main/home.html', {'main_bd': main_bd})

@csrf_exempt
def upload_excel(request):
    global excel_data

    if request.method == 'POST':
        try:
            body = json.loads(request.body)
            excel_data['from_excel'] = body['tables']
            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})

    return JsonResponse({'status': 'error', 'message': 'Invalid request'})

def start_page(request):
    return render(request, 'main/start_page.html')

def excel_popup(request):
    return render(request, 'main/excel_popup.html')

def excel_popup1(request):
    return render(request, 'main/excel_popup1.html')


def get_suggestions(request):
    field = request.GET.get('field')  # Определяем, для какого поля ищем
    query = request.GET.get('query', '').strip()
    direction = request.GET.get('direction', '').strip()  # Получаем введенное направление


    if not query:
        return JsonResponse([], safe=False)  # Если пусто, возвращаем пустой список

    suggestions = []

    if field == 'direction':  # Если запрашивают направление подготовки
        suggestions = list(FirstVariantBd.objects.filter(direction_of_preparation__icontains=query).values_list(
            'direction_of_preparation', flat=True).distinct())
    elif field == 'subject':  # Если запрашивают предмет
        subject_query = FirstVariantBd.objects.filter(name_object__icontains=query)
        if direction:  # Если направление уже введено, фильтруем дополнительно
            subject_query = subject_query.filter(direction_of_preparation__icontains=direction)

        suggestions = list(subject_query.values_list('name_object', flat=True).distinct())

    return JsonResponse(suggestions, safe=False)  # Отправляем данные в JSON

@csrf_exempt
def save_user_data(request):
    """Обработчик для сохранения данных пользователя"""
    if request.method == 'POST':
        # Получаем словарь из сессии или создаем новый, если его нет
        user_data = request.session.get('user_data', {})
        teacher = request.POST.get('teacher', '').strip()
        subject = request.POST.get('subject', '').strip()
        direction = request.POST.get('direction', '').strip()

        user_data['Преподаватель'] = teacher
        user_data['Наименование предмета'] = subject
        user_data['Направление'] = direction
        #print("Сохраненные данные:", user_data)

        obraz_program = FirstVariantBd.objects.filter(name_object=subject,
                                                 direction_of_preparation=direction)  # Фильтрация по предмету

        #print('программа',obraz_program[0].edu_program)
        user_data['Образовательная программа'] = obraz_program[0].edu_program
        request.session['user_data'] = user_data
        request.session.modified = True  # если вносишь изменения в уже существующий словарь

        print("Сохраненные данные:", user_data)
        # for i in obraz_program:
        #     print('программа',i.edu_program)

        return redirect("competencies")

    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})


@csrf_exempt
def save_user_data1(request):
    if request.method == 'POST':
        user_data = request.session.get('user_data', {})
        try:
            data = json.loads(request.body.decode("utf-8"))  # Разбираем JSON из тела запроса

            competencies = data.get("competencies", [])  # Получаем массив компетенций
            user_data["competencies"] = competencies  # Сохраняем их
            request.session['user_data'] = user_data
            request.session.modified = True  # если вносишь изменения в уже существующий словарь

            print("Сохраненные данные:", user_data)  # Выводим для отладки

            return JsonResponse({'status': 'success'})  # Возвращаем успешный ответ

        except json.JSONDecodeError:
            return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)

@csrf_exempt
def save_user_data2(request):
    if request.method == "POST":
        user_data = request.session.get('user_data', {})
        try:
            # Декодируем JSON из запроса
            data = json.loads(request.body)

            # Создаем структуру данных, если ее нет
            if "topics" not in user_data:
                user_data["topics"] = {}

            # Процесс сохранения
            for topic in data.get("topics", []):
                group_name = topic["group"]
                topic_name = topic["topic"]  # Название топика для поиска

                # Если группа уже существует, добавляем новый топик в список этой группы
                if group_name in user_data["topics"]:

                    # Проверяем, существует ли уже такой топик в группе
                    existing_topic = next((t for t in user_data["topics"][group_name] if t["topic"] == topic_name), None)

                    if existing_topic:
                        # Если такой топик уже есть, удаляем его
                        user_data["topics"][group_name].remove(existing_topic)
                        print(f"Топик '{topic_name}' удален из группы '{group_name}'.")

                    user_data["topics"][group_name].append({
                        "topic": topic["topic"],
                        "description": topic["description"],
                        "time": topic["time"]
                    })
                else:
                    # Если группы нет, создаем новую запись с текущим топиком
                    user_data["topics"][group_name] = [{
                        "topic": topic["topic"],
                        "description": topic["description"],
                        "time": topic["time"]
                    }]
            request.session['user_data'] = user_data
            request.session.modified = True  # если вносишь изменения в уже существующий словарь

            print("Данные успешно сохранены:", user_data)  # Логирование

            return JsonResponse({"status": "success"})
        except Exception as e:
            print("Ошибка при обработке данных:", str(e))
            return JsonResponse({"status": "error", "message": str(e)}, status=400)

    return JsonResponse({"status": "invalid request"}, status=400)

@csrf_exempt
def save_user_data3(request):
    if request.method == 'POST':
        user_data = request.session.get('user_data', {})
        curriculum_data = {}
        tables_count = int(request.POST.get("tables_count", 0))

        for i in range(1, tables_count + 1):
            profile_key = f"profile_{i}"
            profiles = request.POST.get(profile_key, f"Профиль {i}")
            topics_data = {}

            row = 1
            while True:
                topic_key = f"topic_{i}_{row}"
                if topic_key not in request.POST:
                    break  # Если нет такой темы — прекращаем цикл по строкам

                topic_name = request.POST.get(topic_key)
                topic_data = {
                    'classroom': request.POST.get(f'classroom_{i}_{row}', ''),
                    'lectures': request.POST.get(f'lectures_{i}_{row}', ''),
                    'seminars': request.POST.get(f'seminars_{i}_{row}', ''),
                    'independent': request.POST.get(f'independent_{i}_{row}', ''),
                }

                topics_data[topic_name] = topic_data
                row += 1

            curriculum_data[f'Профили: {profiles}'] = topics_data

        user_data['curriculum'] = curriculum_data
        request.session['user_data'] = user_data
        request.session.modified = True  # если вносишь изменения в уже существующий словарь

        print(user_data)  # или сохранить куда нужно

        return redirect('content_of_seminars')




@csrf_exempt
def save_user_data4(request):
    if request.method == "POST":
        # Получаем словарь из сессии или создаем новый, если его нет
        user_data = request.session.get('user_data', {})
        try:
            # Загружаем данные из тела запроса
            data = json.loads(request.body)

            # Инициализируем пустой список для сохранения данных
            seminars_content_data = []

            # Проходим по всем переданным данным из таблицы
            for table_data in data.get("tables", []):
                seminars_content = {
                    "topic": table_data.get("topic", ""),
                    "questions": table_data.get("questions", []),
                    "form": table_data.get("form", "")
                }
                seminars_content_data.append(seminars_content)

            # Сохраняем данные в словарь user_data
            user_data["seminars_content"] = seminars_content_data
            request.session['user_data'] = user_data
            request.session.modified = True  # если вносишь изменения в уже существующий словарь

            print("Данные успешно сохранены:", user_data)  # Лог для проверки

            # Возвращаем успешный ответ
            return JsonResponse({"status": "success"})
        except Exception as e:
            print("Ошибка при обработке данных:", str(e))
            return JsonResponse({"status": "error", "message": str(e)}, status=400)

    return JsonResponse({"status": "invalid request"}, status=400)

@csrf_exempt
def save_user_data5(request):
    if request.method == "POST":
        user_data = request.session.get('user_data', {})
        try:
            # Загружаем данные из тела запроса
            data = json.loads(request.body)

            # Инициализируем пустой список для сохранения данных
            questions_list_data = []

            # Проходим по всем переданным данным из таблицы
            for table_data in data.get("tables", []):
                questions_list = {
                    "topic": table_data.get("topic", ""),
                    "questions": table_data.get("questions", []),
                    "form": table_data.get("form", "")
                }
                questions_list_data.append(questions_list)

            # Сохраняем данные в словарь user_data
            user_data["questions_list"] = questions_list_data
            request.session['user_data'] = user_data
            request.session.modified = True  # если вносишь изменения в уже существующий словарь

            print("Данные успешно сохранены:", user_data)  # Лог для проверки

            # Возвращаем успешный ответ
            return JsonResponse({"status": "success"})
        except Exception as e:
            print("Ошибка при обработке данных:", str(e))
            return JsonResponse({"status": "error", "message": str(e)}, status=400)

    return JsonResponse({"status": "invalid request"}, status=400)

@csrf_exempt
def save_user_data6(request):
    if request.method == "POST":
        user_data = request.session.get('user_data', {})
        try:
            data = json.loads(request.body)

            questions_work_data = []
            if "questions" in data:
                questions_work_data = data["questions"]

            # Сохранение в user_data
            user_data["questions_work"] = questions_work_data
            request.session['user_data'] = user_data
            request.session.modified = True  # если вносишь изменения в уже существующий словарь

            print("Данные успешно сохранены:", user_data)  # Проверка данных

            return JsonResponse({"status": "success"})
        except Exception as e:
            print("Ошибка при обработке данных:", str(e))
            return JsonResponse({"status": "error", "message": str(e)}, status=400)

    return JsonResponse({"status": "invalid request"}, status=400)

@csrf_exempt
def save_user_data7(request):
    if request.method == 'POST':
        user_data = request.session.get('user_data', {})
        try:
            data = json.loads(request.body)
            control_tasks = data.get('control_tasks', [])

            if 'control_tasks' not in user_data:
                user_data['control_tasks'] = {}

            for entry in control_tasks:
                profile = entry['profile']
                competence_code = entry['competence_code']
                competence_name = entry['competence_name']
                indicator = entry['indicator']
                know = entry['know']
                do_value = entry['do_value']
                tasks = entry['task']

                if profile not in user_data['control_tasks']:
                    user_data['control_tasks'][profile] = []

                # Проверка на существующую запись
                existing = next((item for item in user_data['control_tasks'][profile]
                                 if item['competence_code'] == competence_code and item['indicator'] == indicator), None)

                if existing:
                    existing['task'].extend(tasks)
                else:
                    user_data['control_tasks'][profile].append({
                        'competence_code': competence_code,
                        'competence_name': competence_name,
                        'indicator': indicator,
                        'know': know,
                        'do': do_value,
                        'task': tasks
                    })
            request.session['user_data'] = user_data
            request.session.modified = True  # если вносишь изменения в уже существующий словарь

            return JsonResponse({'status': 'success'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)

    return JsonResponse({'status': 'error', 'message': 'Invalid method'}, status=405)


@csrf_exempt
def save_user_data8(request):
    if request.method == "POST":
        user_data = request.session.get('user_data', {})
        try:
            data = json.loads(request.body)

            example_quest_data = []
            if "questions" in data:
                example_quest_data = data["questions"]

            # Сохранение в user_data
            user_data["example_quest"] = example_quest_data
            request.session['user_data'] = user_data
            request.session.modified = True  # если вносишь изменения в уже существующий словарь

            print("Данные успешно сохранены:", user_data)  # Проверка данных

            return JsonResponse({"status": "success"})
        except Exception as e:
            print("Ошибка при обработке данных:", str(e))
            return JsonResponse({"status": "error", "message": str(e)}, status=400)

    return JsonResponse({"status": "invalid request"}, status=400)

# def competencies(request):
#     return render(request, 'competencies.html')

def competencies(request):
    user_data = request.session.get('user_data', {})
    subject_name = user_data.get('Наименование предмета', '')  # Извлекаем значение
    print('Название предмета', subject_name)

    direction = user_data.get('Направление', '')
    print('Направление подготовки', direction)
    profiles = FirstVariantBd.objects.filter(name_object=subject_name, direction_of_preparation=direction)  # Фильтрация по предмету

    all_competencies = Description_of_competencies.objects.all()
    competency_dict = {comp.competency_name: comp.description for comp in all_competencies}

    for profile in profiles:
        print(profile.profile)
        profile.competentions_list = []
        for code in profile.competentions.split(','):
            code = code.strip()  # Убираем лишние пробелы
            description = competency_dict.get(code, "Описание не найдено")
            profile.competentions_list.append({'code': code, 'description': description})

    return render(request, 'main/competencies.html', {'profiles': profiles})

def content_of_discipline(request):
    user_data = request.session.get('user_data', {})
    print(user_data)
    if request.method == "POST":
        # Получаем данные из запроса
        data = request.POST.getlist('topic[]')
        descriptions = request.POST.getlist('description[]')
        times = request.POST.getlist('time[]')

        # Сохраняем данные в user_data
        user_data['topics'] = []
        for i in range(len(data)):
            user_data['topics'].append({
                'topic': data[i],
                'description': descriptions[i],
                'time': int(times[i]) if times[i] else 0
            })

        # Отправляем успешный ответ
        return JsonResponse({'status': 'success'})

    # Получаем данные из базы данных
    profiles = FirstVariantBd.objects.filter(
        direction_of_preparation=user_data.get('Направление', ''),
        name_object=user_data.get('Наименование предмета', '')
    )

    # Если профили не найдены, возвращаем ошибку
    if not profiles.exists():
        return render(request, 'main/content_of_discipline.html', {'error': 'Нет данных для отображения'})

    # Группировка профилей по всем важным параметрам
    profile_groups = defaultdict(list)
    profile_data = {}

    for profile in profiles:
        # Создаем уникальный ключ для группировки по всем нужным полям
        key = f"{profile.total_hours}_{profile.classroom_hours}_{profile.lectures}_{profile.seminars}_{profile.independent_work}"

        profile_groups[key].append(profile.profile)
        profile_data[key] = profile  # Сохраняем первый попавшийся объект для этой группы

    # Проверяем, есть ли данные
    print("Profile Groups:", profile_groups)
    print("Profile Data:", profile_data)

    # Передаем данные в шаблон
    return render(request, 'main/content_of_discipline.html', {
        'profile_groups': dict(profile_groups),  # Преобразуем defaultdict в обычный словарь
        'profile_data': profile_data,
        'user_data': user_data  # Передаем user_data, если нужно
    })

def safe_int(value, default=0):
    """Преобразует значение в int, если возможно, иначе возвращает default."""
    try:
        return int(value)
    except (ValueError, TypeError):
        return default


def curriculum(request):
    user_data = request.session.get('user_data', {})
    # Получаем данные из базы
    profiles = FirstVariantBd.objects.filter(
        direction_of_preparation=user_data['Направление'],
        name_object=user_data['Наименование предмета']
    )

    if not profiles.exists():
        return render(request, 'main/curriculum.html', {'error': 'Нет данных для отображения'})

    # Группировка профилей по итоговым значениям
    profile_groups = defaultdict(list)
    profile_data = {}

    for profile in profiles:
        key = (
            profile.total_hours,
            profile.classroom_hours,
            profile.lectures,
            profile.seminars,
            profile.independent_work
        )
        profile_groups[key].append(profile.profile)
        profile_data[key] = profile  # Сохраняем данные профиля

    table_data = []

    for key, profile_list in profile_groups.items():
        summary = profile_data[key]

        # Определение формы контроля
        control_form = "Согласно учебному плану"
        if int(summary.control_work) > 0:
            control_form = "Контрольная работа"
        elif int(summary.essay) > 0:
            control_form = "Эссе"
        elif int(summary.calcul_analytic_work) > 0:
            control_form = "Расчетно-аналитическая работа"
        elif int(summary.creative_homework) > 0:
            control_form = "Домашнее творческое задание"
        elif int(summary.project_work) > 0:
            control_form = "Проектная работа"

        # Проверяем, есть ли профиль в user_data['topics']
        profile_key = ', '.join(profile_list)
        topics = []
        if profile_key in user_data['topics']:
            for topic_info in user_data['topics'][profile_key]:
                topics.append({
                    'topic': topic_info['topic'],
                    'total': topic_info['time'],  # Время из user_data
                    'classroom': safe_int(request.POST.get(f'classroom_{topic_info["topic"]}', 0)),
                    'lectures': safe_int(request.POST.get(f'lectures_{topic_info["topic"]}', 0)),
                    'seminars': safe_int(request.POST.get(f'seminars_{topic_info["topic"]}', 0)),
                    'independent': safe_int(request.POST.get(f'independent_{topic_info["topic"]}', 0)),
                })

        percent_classroom = round((summary.classroom_hours / summary.total_hours) * 100) if summary.total_hours else 0
        percent_lectures = round((summary.lectures / summary.classroom_hours) * 100) if summary.classroom_hours else 0
        percent_seminars = round((summary.seminars / summary.classroom_hours) * 100) if summary.classroom_hours else 0
        percent_independent = round(
            (summary.independent_work / summary.total_hours) * 100) if summary.total_hours else 0

        table_data.append({
            'profiles': profile_key,
            'summary': summary,
            'topics': topics,
            'percent_classroom': percent_classroom,
            'percent_lectures': percent_lectures,
            'percent_seminars': percent_seminars,
            'percent_independent': percent_independent,
            'control_form': control_form
        })
    return render(request, 'main/curriculum.html', {'tables': table_data})

def content_of_seminars(request):
    user_data = request.session.get('user_data', {})
    print(user_data)
    return render(request, "main/content_of_seminars.html", {'user_data': user_data})

def list_of_questions(request):
    user_data = request.session.get('user_data', {})
    print(user_data)
    return render(request, "main/list_of_questions.html", {'user_data': user_data})

def questions_to_work(request):
    user_data = request.session.get('user_data', {})
    print(user_data)
    return render(request, "main/questions_to_work.html")

def example_quest_to_test(request):
    user_data = request.session.get('user_data', {})
    print(user_data)
    return render(request, "main/example_quest_to_test.html")

def example_tasks(request):
    user_data = request.session.get('user_data', {})
    print('проверка', user_data['competencies'])
    competencies_by_profile = {}
    for item in user_data['competencies']:
        profile = item['profile']
        if profile not in competencies_by_profile:
            competencies_by_profile[profile] = []
        competencies_by_profile[profile].append(item)
    print('bhjbcvghh', competencies_by_profile[profile])

    # Добавляем вычисление rowspan для каждой компетенции и объединяем индикаторы, знания и умения
    for profile, competencies in competencies_by_profile.items():
        for competence in competencies:
            competence['rowspan'] = len(competence['indicators']) + 1  # Для каждого индикатора добавляем +1 для самой компетенции
            # Собираем данные для индикаторов, знаний и умений
            competence['indicators_with_know_do'] = list(zip(competence['indicators'], competence['know'], competence['do']))

    return render(request, 'main/example_tasks.html', {'competencies_by_profile': competencies_by_profile})


def insert_table_after(paragraph, rows, cols):
    # Вставляем новый параграф после указанного
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)

    # Получаем объект Paragraph из этого XML
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._element = new_p

    # Создаем таблицу с шириной (например, 6 дюймов)
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.0))

    # Вставляем таблицу сразу после созданного параграфа
    tbl = table._tbl
    new_p.addnext(tbl)

    return table

def insert_paragraph_after(paragraph, text=''):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
    return new_para
def export_to_word(request):
    user_data = request.session.get('user_data', {})
    template_path = 'C:/Users/andru/PycharmProjects/diplommain/main/templates/docx_templates/template_with_placeholders.docx'

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон не найден по пути: {template_path}")

    doc = Document(template_path)

    # Получаем данные из базы
    profiles = FirstVariantBd.objects.filter(
        direction_of_preparation=user_data.get('Направление', ''),
        name_object=user_data.get('Наименование предмета', '')
    )

    if not profiles.exists():
        return HttpResponse("Нет данных для отображения", status=404)

    profile_groups = defaultdict(list)
    profile_data = {}

    for profile in profiles:
        key = f"{profile.total_hours}_{profile.classroom_hours}_{profile.lectures}_{profile.seminars}_{profile.independent_work}"
        profile_groups[key].append(profile.profile)
        profile_data[key] = profile

    def replace_placeholders(paragraphs):
        for p in paragraphs:
            full_text = ''.join(run.text for run in p.runs)

            # Обработка обычных плейсхолдеров
            for key, value in user_data.items():
                if key == 'competencies':
                    continue
                placeholder = f'{{{{{key}}}}}'
                if key in ['Наименование предмета', 'Образовательная программа', 'Направление']:
                    value = f'«{value}»'
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))

            # Обработка таблицы volumes
            if '{{volumes}}' in full_text:
                for run in p.runs:
                    run.text = ''

                for key, profiles_list in profile_groups.items():
                    profile = profile_data[key]

                    # Вставляем абзац с профилями над таблицей — сразу после тега
                    profiles_para = insert_paragraph_after(p, f"Профили: {', '.join(profiles_list)}")

                    sem_list = []
                    for sem_field in [profile.exam, profile.test_obj, profile.test_obj_with_mark]:
                        if sem_field:
                            sem_list += [s.strip() for s in sem_field.split(',') if s.strip().isdigit()]
                    sem_list = sorted(set(sem_list), key=int)

                    columns = 2 + len(sem_list)
                    # Таблица вставляется прямо после текущего параграфа с тегом
                    table = insert_table_after(profiles_para, rows=0, cols=columns)
                    table.style = 'Table Grid'

                    header = table.add_row().cells
                    header[0].text = 'Вид учебной работы по дисциплине'
                    header[1].text = 'Всего в (з/е и часах)'
                    for i, sem in enumerate(sem_list):
                        header[2 + i].text = f'Семестр {sem}'

                    def add_row(name, values):
                        row = table.add_row().cells
                        row[0].text = name
                        for idx, val in enumerate(values):
                            row[1 + idx].text = str(val)

                    ects_total = f"{profile.ECTS}/{profile.total_hours}"
                    add_row('Общая трудоемкость дисциплины', [ects_total] + [profile.total_hours] * len(sem_list))
                    add_row('Контактная работа - Аудиторные занятия', [profile.classroom_hours] * (1 + len(sem_list)))
                    add_row('Лекции', [profile.lectures] * (1 + len(sem_list)))
                    add_row('Семинары, практические занятия', [profile.seminars] * (1 + len(sem_list)))
                    add_row('Самостоятельная работа', [profile.independent_work] * (1 + len(sem_list)))

                    current_control = []
                    if profile.control_work: current_control.append("Контрольная работа")
                    if profile.essay: current_control.append("Реферат")
                    if profile.calcul_analytic_work: current_control.append("Расчетно-аналитическая работа")
                    if profile.creative_homework: current_control.append("Творческое задание")
                    if profile.project_work: current_control.append("Проектная работа")
                    add_row('Вид текущего контроля', [", ".join(current_control)] + current_control[:len(sem_list)])

                    attestation = []
                    if profile.exam: attestation.append("Экзамен")
                    if profile.test_obj: attestation.append("Зачет")
                    if profile.test_obj_with_mark: attestation.append("Зачет с оценкой")
                    add_row('Вид промежуточной аттестации', [", ".join(attestation)] + attestation[:len(sem_list)])

            elif full_text != ''.join(run.text for run in p.runs):
                    for run in p.runs:
                        run.text = ''
                    p.add_run(full_text)


            elif '{{competencies}}' in full_text:

                for run in p.runs:
                    run.text = run.text.replace('{{competencies}}', '')

                table = insert_table_after(p, rows=1, cols=4)

                table.style = 'Table Grid'

                # Заголовки

                headers = ['Код компетенции', 'Наименование компетенции', 'Индикаторы достижения компетенции',
                           'Результаты обучения']

                hdr_cells = table.rows[0].cells

                for i, text in enumerate(headers):
                    hdr_cells[i].text = text

                    para = hdr_cells[i].paragraphs[0]

                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = para.runs[0]

                    run.font.bold = True

                    run.font.size = Pt(11)

                # Группировка по профилям

                grouped = defaultdict(list)

                for c in user_data['competencies']:
                    grouped[c['profile']].append(c)

                for profile, competencies in grouped.items():

                    # Строка профиля

                    profile_row = table.add_row().cells

                    profile_cell = profile_row[0]

                    profile_cell.text = profile

                    for i in range(1, 4):
                        profile_row[i].text = ''

                    profile_row[0].merge(profile_row[1]).merge(profile_row[2]).merge(profile_row[3])

                    para = profile_cell.paragraphs[0]

                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    para.runs[0].bold = True

                    para.runs[0].font.size = Pt(12)

                    # Компетенции

                    for comp in competencies:

                        indicators = comp['indicators']

                        knows = comp['know']

                        dos = comp['do']

                        num_rows = len(indicators)

                        for idx in range(num_rows):

                            row = table.add_row().cells

                            # Удаляем переносы строк, если есть

                            code = comp['competence_code'].replace('\n', ' ').strip()

                            name = comp['competence_name'].replace('\n', ' ').strip()

                            # Только первая строка — с кодом и названием

                            if idx == 0:

                                row[0].text = code

                                row[1].text = name

                            else:

                                row[0].text = ''

                                row[1].text = ''

                            # Индикатор

                            row[2].text = indicators[idx]

                            # Результаты обучения по индексу

                            know_text = knows[idx] if idx < len(knows) else ''

                            do_text = dos[idx] if idx < len(dos) else ''

                            row[3].text = f"Знать:\n{know_text}\nУметь:\n{do_text}"

                            # Выравнивание

                            for cell in row:
                                para = cell.paragraphs[0]

                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif '{{content}}' in full_text:
                # Удаляем текст метки
                for run in p.runs:
                    run.text = ''

                # Получаем первую группу тем
                first_group_key = next(iter(user_data.get('topics', {})), None)
                if first_group_key:
                    topics_list = user_data['topics'][first_group_key]

                    for item in topics_list:
                        topic_para = insert_paragraph_after(p, '')
                        topic_run = topic_para.add_run(item['topic'])
                        topic_run.bold = True
                        topic_para.paragraph_format.first_line_indent = Inches(0.5)

                        desc_para = insert_paragraph_after(topic_para, item['description'])
                        desc_para.paragraph_format.first_line_indent = Inches(0.5)

            # Обработка плейсхолдера {{curriculum}}
            elif '{{syllabus}}' in full_text:
                for run in p.runs:
                    run.text = ''

                topics_data = user_data.get("topics", {})
                curriculum_data = user_data.get("curriculum", {})

                for key, profile_list in profile_groups.items():
                    profile_key = ', '.join(profile_list)
                    topics = topics_data.get(profile_key, [])
                    summary = profile_data.get(key)
                    profile_curriculum = curriculum_data.get(f"Профили: {profile_key}", {})

                    # Название профиля
                    paragraph_before = p.insert_paragraph_before(f"Профиль: {profile_key}", style='Normal')

                    table = insert_table_after(paragraph_before, rows=1, cols=8)
                    table.style = 'Table Grid'

                    # Заголовки таблицы
                    hdr1 = table.rows[0].cells
                    hdr1[0].text = '№ п/п'
                    hdr1[1].text = 'Наименование тем (разделов) дисциплины'
                    hdr1[2].text = 'Всего'
                    hdr1[3].text = 'Контактная работа * - Аудиторная работа'
                    hdr1[6].text = 'Самостоятельная работа'
                    hdr1[7].text = 'Формы текущего контроля успеваемости'
                    hdr1[3].merge(hdr1[4]).merge(hdr1[5])

                    row = table.add_row().cells
                    row[3].text = 'Общая, в т.ч.:'
                    row[4].text = 'Лекции'
                    row[5].text = 'Семинары, практические занятия'

                    # Заполнение тем
                    for i, topic_info in enumerate(topics, 1):
                        topic_name = topic_info.get('topic', '')
                        topic_hours = profile_curriculum.get(topic_name, {})

                        row = table.add_row().cells
                        row[0].text = str(i)
                        row[1].text = topic_name
                        row[2].text = str(int(topic_hours.get('classroom', ''))+int(topic_hours.get('independent', '')))
                        row[3].text = topic_hours.get('classroom', '')
                        row[4].text = topic_hours.get('lectures', '')
                        row[5].text = topic_hours.get('seminars', '')
                        row[6].text = topic_hours.get('independent', '')
                        row[7].text = 'Дискуссия, защита практических заданий'

                    # Форма контроля
                    control_form = "Согласно учебному плану: "
                    if int(summary.control_work) > 0:
                        control_form += "Контрольная работа"
                    elif int(summary.essay) > 0:
                        control_form += "Эссе"
                    elif int(summary.calcul_analytic_work) > 0:
                        control_form += "Расчетно-аналитическая работа"
                    elif int(summary.creative_homework) > 0:
                        control_form += "Домашнее творческое задание"
                    elif int(summary.project_work) > 0:
                        control_form += "Проектная работа"

                    # Строка "В целом по дисциплине"
                    row = table.add_row().cells
                    row[0].merge(row[1])
                    row[0].text = 'В целом по дисциплине'
                    row[2].text = str(summary.total_hours)
                    row[3].text = str(summary.classroom_hours)
                    row[4].text = str(summary.lectures)
                    row[5].text = str(summary.seminars)
                    row[6].text = str(summary.independent_work)
                    row[7].text = control_form

                    # Строка "Итого в %"
                    row = table.add_row().cells
                    row[0].merge(row[1])
                    row[0].text = 'Итого в %'
                    row[2].text = '100'
                    row[3].text = str(
                        round((summary.classroom_hours / summary.total_hours) * 100)) if summary.total_hours else '0'
                    row[4].text = str(
                        round((summary.lectures / summary.classroom_hours) * 100)) if summary.classroom_hours else '0'
                    row[5].text = str(
                        round((summary.seminars / summary.classroom_hours) * 100)) if summary.classroom_hours else '0'
                    row[6].text = str(
                        round((summary.independent_work / summary.total_hours) * 100)) if summary.total_hours else '0'
                    row[7].text = ''

            elif '{{content_for_seminars}}' in full_text:
                for run in p.runs:
                    run.text = ''

                seminars_content = user_data.get("seminars_content", [])

                # Вставляем таблицу после текущего абзаца
                table = insert_table_after(p, rows=1, cols=3)
                table.style = 'Table Grid'

                # Заголовки таблицы
                hdr = table.rows[0].cells
                hdr[0].text = 'Наименование тем (разделов) дисциплины'
                hdr[1].text = 'Перечень вопросов, отводимых на самостоятельное освоение'
                hdr[2].text = 'Формы внеаудиторной самостоятельной работы'

                # Заполнение строк таблицы
                for item in seminars_content:
                    row = table.add_row().cells
                    row[0].text = item.get("topic", "")
                    row[1].text = '\n'.join(f'– {q}' for q in item.get("questions", []))
                    row[2].text = item.get("form", "")

            elif '{{list_questions}}' in full_text:
                for run in p.runs:
                    run.text = ''

                seminars_content = user_data.get("questions_list", [])

                # Вставляем таблицу после текущего абзаца
                table = insert_table_after(p, rows=1, cols=3)
                table.style = 'Table Grid'

                # Заголовки таблицы
                hdr = table.rows[0].cells
                hdr[0].text = 'Наименование тем (разделов) дисциплины'
                hdr[1].text = 'Перечень вопросов, отводимых на самостоятельное освоение'
                hdr[2].text = 'Формы внеаудиторной самостоятельной работы'

                # Заполнение строк таблицы
                for item in seminars_content:
                    row = table.add_row().cells
                    row[0].text = item.get("topic", "")
                    row[1].text = '\n'.join(f'– {q}' for q in item.get("questions", []))
                    row[2].text = item.get("form", "")

            elif '{{questions_to_control}}' in full_text:
                for run in p.runs:
                    run.text = ''

                questions = user_data.get('questions_work', [])

                for i, question in enumerate(questions, 1):
                    new_paragraph = p.insert_paragraph_before(f"{i}. {question}", style='Normal')

            elif '{{tasks_for_control}}' in full_text:

                for run in p.runs:
                    run.text = ''

                table = insert_table_after(p, rows=1, cols=4)
                table.style = 'Table Grid'

                # Заголовки
                headers = ['Компетенция', 'Индикатор', 'Результаты обучения', 'Типовые контрольные задания']
                hdr_cells = table.rows[0].cells

                for i, text in enumerate(headers):
                    hdr_cells[i].text = text
                    para = hdr_cells[i].paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(11)

                # Группировка по профилям и компетенциям
                grouped = defaultdict(lambda: defaultdict(list))

                for profile, items in user_data['control_tasks'].items():
                    for item in items:
                        comp_key = (item['competence_code'], item['competence_name'])
                        grouped[profile][comp_key].append(item)

                # Заполнение таблицы
                for profile, competencies in grouped.items():
                    # Строка с названием профиля
                    profile_row = table.add_row().cells
                    profile_cell = profile_row[0]
                    profile_cell.text = profile

                    for i in range(1, 4):
                        profile_row[i].text = ''

                    profile_row[0].merge(profile_row[1]).merge(profile_row[2]).merge(profile_row[3])
                    para = profile_cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.runs[0].bold = True
                    para.runs[0].font.size = Pt(12)

                    for (code, name), items in competencies.items():
                        for idx, item in enumerate(items):
                            row = table.add_row().cells

                            # Компетенция: только в первой строке
                            if idx == 0:
                                row[0].text = f"{code}\n{name}"
                            else:
                                row[0].text = ''

                            # Индикатор
                            row[1].text = item['indicator']

                            # Результаты обучения
                            know = item.get('know', '')
                            do = item.get('do', '')
                            row[2].text = f"Знать:\n{know}\nУметь:\n{do}"

                            # Типовые задания
                            tasks = item.get('task', [])
                            if isinstance(tasks, list):
                                row[3].text = '\n'.join(f'Задание {i+1}. {task}' for i, task in enumerate(tasks) if task.strip())
                            else:
                                row[3].text = str(tasks)

                            #print('вывод', row[3].text)

                            # Выравнивание по центру
                            for cell in row:
                                para = cell.paragraphs[0]
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif '{{questions_to_test_obj}}' in full_text:
                for run in p.runs:
                    run.text = ''

                questions = user_data.get('example_quest', [])

                for i, question in enumerate(questions, 1):
                    new_paragraph = p.insert_paragraph_before(f"{i}. {question}", style='Normal')

    # Обработка параграфов
    replace_placeholders(doc.paragraphs)

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders(cell.paragraphs)

    # Отправка файла
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=program.docx'
    doc.save(response)
    return response